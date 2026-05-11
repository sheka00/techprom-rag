import logging
import warnings
from pathlib import Path
from docx import Document
import shutil
from docling.document_converter import DocumentConverter

# Игнорируем предупреждения от docling
warnings.filterwarnings("ignore")
logging.getLogger("docling").setLevel(logging.ERROR)
logging.getLogger("docling.backend").setLevel(logging.ERROR)

logger = logging.getLogger(__name__)


class DocxPreprocessor:
    def __init__(self, output_dir: str = "data/preprocessed"):
        self.output_dir = Path(output_dir)
        self.converter = DocumentConverter()
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def process_file(self, input_path: Path) -> Path:
        """
        Обрабатывает один docx файл: добавляет пустые строки после заголовков,
        сравнивает результат конвертации через docling и сохраняет лучший вариант
        в папку вывода.
        """
        logger.info(f"🔄 Предобработка {input_path.name}...")

        # Чтение документа
        doc = Document(str(input_path))

        # Временное имя для модифицированного файла
        temp_modified_path = self.output_dir / f"temp_{input_path.name}"

        # Находим все заголовки и их позиции
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name.startswith("Heading") and paragraph.text.strip():
                headings.append((i, paragraph))

        # Добавляем пустые строки там, где их нет после заголовков
        added_count = 0
        # Идем с конца, чтобы не нарушить индексацию
        for index, heading_para in reversed(headings):
            try:
                # Проверяем, есть ли пустая строка после заголовка
                next_index = index + 1
                need_to_add = True

                # Проверяем, не последний ли это параграф
                if next_index < len(doc.paragraphs):
                    next_para = doc.paragraphs[next_index]
                    # Если следующий параграф пустой или содержит только пробелы
                    if next_para.text.strip() == "":
                        need_to_add = False

                if need_to_add:
                    # Вставляем пустой параграф ПОСЛЕ заголовка
                    if next_index < len(doc.paragraphs):
                        doc.paragraphs[next_index].insert_paragraph_before("")
                    else:
                        doc.add_paragraph("")
                    added_count += 1

            except Exception as e:
                logger.debug(
                    f"Ошибка при обработке параграфа {index} в {input_path.name}: {e}"
                )

        # Сохранение модифицированного документа
        doc.save(str(temp_modified_path))

        try:
            # Конвертируем исходный файл
            result_orig = self.converter.convert(str(input_path))
            original_text = result_orig.document.export_to_markdown()
            original_length = len(original_text)

            # Конвертируем измененный файл
            result_mod = self.converter.convert(str(temp_modified_path))
            modified_text = result_mod.document.export_to_markdown()
            modified_length = len(modified_text)

            difference = modified_length - original_length
            final_path = self.output_dir / input_path.name

            if difference > 0:
                logger.info(
                    f"   ✅ Улучшение: +{difference} симв. (добавлено строк: {added_count})"
                )
                if final_path.exists():
                    final_path.unlink()
                temp_modified_path.rename(final_path)
                return final_path
            else:
                logger.info(
                    f"   ℹ️  Без изменений (разница {difference} симв.). Используем оригинал."
                )
                # Удаляем временный файл
                if temp_modified_path.exists():
                    temp_modified_path.unlink()

                # Копируем или линкуем оригинал в папку вывода
                # Для простоты - копируем
                shutil.copy2(input_path, final_path)
                return final_path

        except Exception as e:
            logger.error(
                f"   ❌ Ошибка при сравнении вариантов для {input_path.name}: {e}"
            )
            # В случае ошибки просто копируем оригинал
            final_path = self.output_dir / input_path.name
            shutil.copy2(input_path, final_path)
            if temp_modified_path.exists():
                temp_modified_path.unlink()
            return final_path

    def process_directory(self, input_dir: str):
        input_path = Path(input_dir)
        files = list(input_path.rglob("*.docx"))
        logger.info(f"📂 Найдено файлов для предобработки: {len(files)}")

        processed_files = []
        for f in files:
            # Сохраняем структуру подпапок если нужно?
            # Пользователь просил "создай новую папку и туда сохраняй файлы"
            # Если в data/ТИ есть подпапки, сохраним их структуру
            rel_path = f.relative_to(input_path)
            target_file_path = self.output_dir / rel_path
            target_file_path.parent.mkdir(parents=True, exist_ok=True)

            # Переопределяем process_file чтобы работало с относительным путем
            self._process_single_file(f, target_file_path)
            processed_files.append(target_file_path)

        return processed_files

    def _process_single_file(self, input_path: Path, output_path: Path):
        """Внутренний метод для процессинга с явным путем назначения"""
        doc = Document(str(input_path))
        temp_modified_path = output_path.with_name(f"temp_{output_path.name}")
        temp_modified_path.parent.mkdir(parents=True, exist_ok=True)

        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style.name.startswith("Heading") and paragraph.text.strip():
                headings.append((i, paragraph))

        added_count = 0
        for index, heading_para in reversed(headings):
            try:
                next_index = index + 1
                need_to_add = True
                if next_index < len(doc.paragraphs):
                    next_para = doc.paragraphs[next_index]
                    if next_para.text.strip() == "":
                        need_to_add = False
                if need_to_add:
                    if next_index < len(doc.paragraphs):
                        doc.paragraphs[next_index].insert_paragraph_before("")
                    else:
                        doc.add_paragraph("")
                    added_count += 1
            except Exception:
                pass

        doc.save(str(temp_modified_path))

        try:
            result_orig = self.converter.convert(str(input_path))
            original_text = result_orig.document.export_to_markdown()
            original_length = len(original_text)

            result_mod = self.converter.convert(str(temp_modified_path))
            modified_text = result_mod.document.export_to_markdown()
            modified_length = len(modified_text)

            if modified_length > original_length:
                logger.info(
                    f"   ✅ {input_path.name}: +{modified_length - original_length} симв."
                )
                if output_path.exists():
                    output_path.unlink()
                temp_modified_path.rename(output_path)
            else:
                logger.debug(f"   ℹ️  {input_path.name}: без улучшений.")
                if temp_modified_path.exists():
                    temp_modified_path.unlink()
                shutil.copy2(input_path, output_path)
        except Exception as e:
            logger.error(f"   ❌ Ошибка {input_path.name}: {e}")
            shutil.copy2(input_path, output_path)
            if temp_modified_path.exists():
                temp_modified_path.unlink()
