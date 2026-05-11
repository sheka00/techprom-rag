import os
import json
import logging
import re
import requests
import time
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
from tqdm import tqdm
import argparse

# Настройка логирования для бенчмарка
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger("Benchmark")


class RAGBenchmark:
    def __init__(
        self,
        test_docx_path: str = "data/тестовые_вопросы.docx",
        api_url: str = "http://127.0.0.1:8000",
        ollama_url: str = "http://ollama:11434",
        report_path: str = "data/benchmark_report.json",
    ):
        self.api_url = api_url.rstrip("/")
        self.ollama_url = ollama_url.rstrip("/")
        self.test_data_path = test_docx_path
        self.report_path = Path(report_path)

        # Проверяем доступность API
        try:
            resp = requests.get(f"{self.api_url}/health", timeout=5)
            if resp.status_code == 200:
                logger.info(f"✅ API доступно по адресу: {self.api_url}")
            else:
                logger.warning(
                    f"⚠️ API вернуло статус {resp.status_code} по адресу: {self.api_url}/health"
                )
        except Exception as e:
            logger.error(f"❌ API недоступно по адресу {self.api_url}: {e}")
            logger.info("Убедитесь, что сервис 'api' запущен и слушает порт 8000.")

    def parse_test_data(self) -> List[Dict[str, Any]]:
        """Парсит DOCX с тестовыми вопросами"""
        if not os.path.exists(self.test_data_path):
            raise FileNotFoundError(f"Файл не найден: {self.test_data_path}")

        doc = Document(self.test_data_path)

        test_cases = []
        if not doc.tables:
            logger.warning(
                f"⚠️ В файле {self.test_data_path} не найдено таблиц. Пытаюсь парсить текст..."
            )
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            current_q = None
            current_a = []

            for p in paragraphs:
                # Если параграф заканчивается на '?' - это вероятно новый вопрос
                if p.endswith("?"):
                    if current_q:
                        # Сохраняем предыдущую найденную пару
                        test_cases.append(
                            {
                                "id": str(len(test_cases) + 1),
                                "expected_doc": "DOC",
                                "question": current_q,
                                "expected_answer": "\n".join(current_a)
                                .replace("Ответ:", "")
                                .strip(),
                                "expected_source": "",
                            }
                        )
                    current_q = p
                    current_a = []
                elif current_q:
                    current_a.append(p)

            # Не забываем добавить последний вопрос и ответ
            if current_q:
                test_cases.append(
                    {
                        "id": str(len(test_cases) + 1),
                        "expected_doc": "DOC",
                        "question": current_q,
                        "expected_answer": "\n".join(current_a)
                        .replace("Ответ:", "")
                        .strip(),
                        "expected_source": "",
                    }
                )

            if not test_cases:
                logger.error(
                    "❌ Не удалось найти вопросы (заканчивающиеся на '?') в тексте файла."
                )
                return []
        else:
            table = doc.tables[0]
            for i, row in enumerate(table.rows):
                if i == 0:
                    continue  # Пропускаем заголовок

                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) < 4:
                    continue

                # Структура: [№, Doc Name, Question, Answer, Source]
                doc_full_name = cells[1]
                question = cells[2]
                expected_answer = cells[3]
                expected_source = cells[4] if len(cells) > 4 else ""

                # Извлекаем код документа (например "ТИ 001" или "СТО 01" или "ТП.И06")
                match = re.search(r"(?:ТИ|СТО|ТП\.И)\s*\d+", doc_full_name)
                expected_doc = match.group(0) if match else doc_full_name

                # Дополнительная очистка для СТО (может быть СТО 01-2022)
                if "СТО" in expected_doc:
                    match_sto = re.search(r"СТО\s*\d+-\d+", doc_full_name)
                    if match_sto:
                        expected_doc = match_sto.group(0)

                if question and expected_answer:
                    test_cases.append(
                        {
                            "id": cells[0],
                            "expected_doc": expected_doc,
                            "question": question,
                            "expected_answer": expected_answer,
                            "expected_source": expected_source,
                        }
                    )

        logger.info(f"Загружено {len(test_cases)} тестовых вопросов")
        return test_cases

    def evaluate_answer_with_llm(
        self,
        question: str,
        expected_ans: str,
        expected_src: str,
        actual_ans: str,
        actual_src: str,
        actual_quote: str,
    ) -> bool:
        """LLM-as-a-judge: проверка правильности ответа через прямой запрос к Ollama"""
        prompt = f"""Ниже приведен технический вопрос, ЭТАЛОННЫЙ ответ (Ground Truth) и ответ от RAG-системы.
Твоя задача — выступить в роли СПРАВЕДЛИВОГО технического эксперта и определить, является ли ответ RAG-системы верным по смыслу.

ВОПРОС: {question}

--- ЭТАЛОН (GROUND TRUTH) ---
ОТВЕТ: {expected_ans}
ИСТОЧНИК: {expected_src}

--- ОТВЕТ RAG-СИСТЕМЫ ---
ОТВЕТ: {actual_ans}
ИСТОЧНИК: {actual_src}
ЦИТАТА: {actual_quote}

КРИТЕРИИ ОЦЕНКИ:
1. ЧИСЛОВАЯ И ФАКТИЧЕСКАЯ ТОЧНОСТЬ (ГЛАВНОЕ): Любые числовые значения (размеры, время, температура) должны СТРОГО совпадать. Если в эталоне "0,3", а в ответе "0,35" — это ГРУБАЯ ОШИБКА (ИТОГ: НЕТ). Замена конкретных объектов на совершенно другие (например, "К-17" на "любое масло" или "контрольное расстояние" на "контроль соответствия размеров") — это НЕВЕРНО.
2. ПРОГРАММНЫЙ МУСОР: Если ответ содержит фрагменты JSON-синтаксиса (кавычки, двоеточия, ключи словарей вроде "temperature": 100 или 1000":"#1) — это расценивается как галлюцинация формата и оценивается как НЕВЕРНО, даже если там есть правильная цифра. Допустим только формат простого списка значений.
3. СИНОНИМЫ И ТЕРМИНОЛОГИЯ: Очевидные и общепринятые технические синонимы считаются абсолютно верными! Например, "термообработка" и "термическая обработка" — это ОДНО И ТО ЖЕ (ИТОГ: ДА). Игнорируй разницу в полных и кратких формах терминов.
4. ФОРМАТ: Игнорируй точки в конце, регистр букв и порядок слов, если суть полностью идентична и числа совпадают.
5. ОТКАЗ: Если ответ "Информация не найдена", а в эталоне ответ есть — это НЕВЕРНО.

ЛОЯЛЬНОСТЬ: Будь МАКСИМАЛЬНО лоялен к очевидным синонимам (термообработка = термическая обработка), но АБСОЛЮТНО СТРОГ к цифрам (0.3 != 0.35), программному мусору (наличию кусков JSON вроде "key": value) и подмене понятий (одно действие или предмет вместо другого).

Сначала напиши КРАТКОЕ обоснование (1 предложение), а затем на новой строке напиши строго 'ИТОГ: ДА' (если верно) или 'ИТОГ: НЕТ' (если неверно).
"""

        payload = {
            "model": os.getenv(
                "OLLAMA_MODEL", "qwen3:14b"
            ),  # Используем текущую модель
            "messages": [{"role": "user", "content": prompt}],
            "stream": False,
            "think": False,
            "options": {
                "temperature": 0.0,
                "num_ctx": int(os.getenv("OLLAMA_NUM_CTX", "16000")),
                "seed": 42,
            },
        }

        max_retries = 3
        ollama_timeout = 300

        for attempt in range(max_retries):
            try:
                response = requests.post(
                    f"{self.ollama_url}/api/chat", json=payload, timeout=ollama_timeout
                )
                response.raise_for_status()
                text = (
                    response.json()
                    .get("message", {})
                    .get("content", "")
                    .strip()
                    .upper()
                )

                # Ищем "ИТОГ: ДА" или "ИТОГ: НЕТ"
                if "ИТОГ: ДА" in text:
                    return True
                if "ИТОГ: НЕТ" in text:
                    return False

                # Если формат нарушен, но ДА/НЕТ в тексте есть (крайний случай)
                if "ИТОГ" not in text:
                    if text.endswith("ДА"):
                        return True
                    if text.endswith("НЕТ"):
                        return False

                return False
            except Exception:
                if attempt < max_retries - 1:
                    time.sleep(2)
                else:
                    return False
        return False

    def _save_report(
        self,
        results: List[Dict],
        retrieval_hits: int,
        correct_answers: int,
        total_expected: int,
    ):
        processed = len(results)

        report_data = {
            "metrics": {
                "total_expected": total_expected,
                "processed": processed,
                "retrieval_accuracy": retrieval_hits / processed
                if processed > 0
                else 0,
                "answer_quality": correct_answers / processed if processed > 0 else 0,
                "retrieval_hits": retrieval_hits,
                "correct_answers": correct_answers,
            },
            "details": results,
        }

        try:
            self.report_path.parent.mkdir(parents=True, exist_ok=True)
            temp_path = self.report_path.with_suffix(".json.tmp")
            with open(temp_path, "w", encoding="utf-8") as f:
                json.dump(report_data, f, ensure_ascii=False, indent=2)
            temp_path.replace(self.report_path)
        except Exception as e:
            logger.error(f"Ошибка при сохранении отчета: {e}")

    def run(
        self, limit: int = None, collection_name: str = None, system_prompt: str = None
    ):
        test_cases = self.parse_test_data()
        if limit:
            test_cases = test_cases[:limit]
            logger.info(f"Ограничение бенчмарка: только первые {limit} вопросов")

        results = []
        retrieval_hits = 0
        correct_answers = 0

        print("\n" + "=" * 80)
        print(f"{'№':<3} | {'Doc':<10} | {'Retr':<4} | {'LLM':<4} | {'Question'}")
        print("-" * 80)

        for case in tqdm(test_cases, desc="Бенчмарк"):
            try:
                # 1. Запрос к RAG через API
                try:
                    query_start = time.time()
                    payload = {"question": case["question"]}
                    if collection_name:
                        payload["collection_name"] = collection_name
                    if system_prompt:
                        payload["system_prompt"] = system_prompt

                    resp = requests.post(
                        f"{self.api_url}/query",
                        json=payload,
                        timeout=600,  # Таймаут с запасом для RAG API
                    )
                    resp.raise_for_status()
                    query_result = resp.json()
                    actual_answer = query_result["answer"]
                    query_time = time.time() - query_start
                    logger.debug(f"Запрос выполнен за {query_time:.1f} сек.")
                except Exception as e:
                    logger.error(f"❌ Ошибка API на вопросе {case['id']}: {e}")
                    actual_answer = f"ОШИБКА API: {e}"
                    query_result = {"sources": [], "source": "", "quote": ""}

                # 2. Проверка поиска (Retrieval)
                is_retrieval_ok = False
                expected_doc_clean = case["expected_doc"].lower().replace(" ", "")
                for src in query_result.get("sources", []):
                    doc_name = src.get("document_name", "").lower().replace(" ", "")
                    # Проверяем вхождение кода (например СТО 01 или ТП.И06 в названии файла)
                    if expected_doc_clean in doc_name:
                        is_retrieval_ok = True
                        break

                if is_retrieval_ok:
                    retrieval_hits += 1

                # 3. Проверка качества ответа (LLM judge)
                is_answer_ok = False
                if "ОШИБКА" not in actual_answer:
                    is_answer_ok = self.evaluate_answer_with_llm(
                        question=case["question"],
                        expected_ans=case["expected_answer"],
                        expected_src=case["expected_source"],
                        actual_ans=actual_answer,
                        actual_src=query_result.get("source", ""),
                        actual_quote=query_result.get("quote", ""),
                    )

                if is_answer_ok:
                    correct_answers += 1

                results.append(
                    {
                        "case": case,
                        "retrieval_ok": is_retrieval_ok,
                        "answer_ok": is_answer_ok,
                        "actual_answer": actual_answer,
                        "sources": [
                            s.get("document_name", "")
                            for s in query_result.get("sources", [])
                        ],
                    }
                )

                # Вывод строки результата
                print(
                    f"{case['id']:<3} | {case['expected_doc']:<10} | {'✅' if is_retrieval_ok else '❌':<4} | {'✅' if is_answer_ok else '❌':<4} | {case['question'][:50]}..."
                )

                # Сохраняем промежуточные результаты
                self._save_report(
                    results, retrieval_hits, correct_answers, len(test_cases)
                )

            except Exception as e:
                logger.error(f"Ошибка при обработке вопроса {case['id']}: {e}")

        # Итоговая статистика
        total = len(test_cases)
        print("\n" + "=" * 80)
        print("📊 ИТОГО:")
        print(
            f"✅ Retrieval Accuracy: {retrieval_hits / total:.2%} ({retrieval_hits}/{total})"
        )
        print(
            f"✅ LLM Answer Quality: {correct_answers / total:.2%} ({correct_answers}/{total})"
        )
        print("=" * 80)

        self._save_report(results, retrieval_hits, correct_answers, total)
        logger.info(f"Детальный отчет обновлен в: {self.report_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Запуск бенчмарка RAG через API")
    parser.add_argument(
        "--docx",
        type=str,
        default="data/тестовые_вопросы.docx",
        help="Путь к DOCX с тестами",
    )
    parser.add_argument(
        "--api-url", type=str, default="http://127.0.0.1:8000", help="URL API сервиса"
    )
    parser.add_argument(
        "--ollama-url",
        type=str,
        default="http://ollama:11434",
        help="URL Ollama сервиса",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Ограничить количество тестовых вопросов",
    )
    parser.add_argument(
        "--collection", type=str, default=None, help="Имя коллекции для запросов"
    )
    parser.add_argument(
        "--report",
        type=str,
        default="data/benchmark_report.json",
        help="Путь для сохранения отчета",
    )
    parser.add_argument(
        "--system-prompt",
        type=str,
        default=None,
        help="Переопределить системный промпт для всех запросов",
    )
    parser.add_argument(
        "--detailed",
        action="store_true",
        help="Использовать усиленный промпт для развернутых ответов",
    )
    args = parser.parse_args()

    custom_prompt = args.system_prompt
    if args.detailed:
        custom_prompt = """Ты — эксперт-аналитик технической документации. 
Твоя задача: давать максимально подробные, развернутые и точные ответы, основываясь ТОЛЬКО на предоставленных данных.

ОТВЕЧАЙ СТРОГО В ФОРМАТЕ JSON С КЛЮЧАМИ:
{
  "answer": "развернутый и технически точный ответ на вопрос",
  "source": "название документа и конкретный раздел/пункт",
  "quote": "точная цитата из текста, подтверждающая наиболее важную часть ответа"
}

ПРАВИЛА:
1. Если информации в контексте недостаточно — в ключе "answer" отвечай "Информация не найдена".
2. Отвечай технически точно, ПОДРОБНО и по существу. Обязательно учитывай все детали из предоставленных фрагментов: технические характеристики, условия, шаги процессов, ограничения и исключения.
3. Если описание процесса содержит важные условия или разные способы действия (например, для разных типов резьб или деталей), обязательно сохрани эту структуру и опиши все детали (используй маркированные списки). 
4. Не сокращай ответ. Если в данных есть описание "как", "почему" или "при каких условиях", включи это в ответ.
5. Избегай вводных слов и общих фраз. Сразу переходи к сути дела.
6. Используй ТОЛЬКО ключи "answer", "source", "quote". Ключи должны быть на английском.
7. ОТВЕЧАЙ ТОЛЬКО НА РУССКОМ ЯЗЫКЕ."""

    benchmark = RAGBenchmark(
        test_docx_path=args.docx,
        api_url=args.api_url,
        ollama_url=args.ollama_url,
        report_path=args.report,
    )
    benchmark.run(
        limit=args.limit, collection_name=args.collection, system_prompt=custom_prompt
    )
