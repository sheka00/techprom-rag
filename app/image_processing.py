import logging
import os
import re
import subprocess
import tempfile
import zipfile
import base64
from pathlib import Path
from typing import Callable, List, Optional, Tuple

import fitz  # PyMuPDF
import pypandoc
from docx import Document as DocxDocument
from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.datamodel.base_models import InputFormat
from docling_core.types.doc.document import TableItem
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI

from .document_parser import DocumentChunk

logger = logging.getLogger(__name__)

# ─── Константы и настройки ───────────────────────────────────────────────────

_SKIP_IMG_EXTS = {".emf", ".wmf", ".svg", ".wdp"}
_MIME = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".webp": "image/webp",
}
_CTX_WINDOW = 300
_TABLE_MAX_CHARS = 6000
_DUP_CAPTION_MARKER = "см. описание выше"

# ─── Helpers: Кадрирование и OCR ─────────────────────────────────────────────


def is_garbage(
    text: str,
    threshold: float = 0.30,
    num_pages: Optional[int] = None,
    min_chars_per_page: int = 100,
) -> bool:
    tokens = [t for t in re.split(r"\s+", text.strip()) if t]
    if not tokens:
        return True
    if num_pages and len(text) / num_pages < min_chars_per_page:
        return True
    return sum(1 for t in tokens if len(t) == 1) / len(tokens) > threshold


def _strip_image_attrs(md: str) -> str:
    return re.sub(r"\]\(([^)]+)\)\{[^}]+\}", r"](\1)", md)


def _extract_context(md: str, pos: int, end: int) -> str:
    before = md[max(0, pos - _CTX_WINDOW) : pos].strip()
    after = md[end : end + _CTX_WINDOW].strip()
    before = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", before).strip()
    after = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", after).strip()
    parts = [p for p in [before, after] if p]
    return "\n...\n".join(parts)


# ─── Таблицы ─────────────────────────────────────────────────────────────────


def _html_table_to_md(content: str) -> str:
    m = re.search(r"<table\b", content, re.IGNORECASE)
    if not m:
        return content
    prefix = content[: m.start()].strip()
    table_html = content[m.start() :]
    rows = re.findall(r"<tr\b[^>]*>(.*?)</tr>", table_html, re.DOTALL | re.IGNORECASE)
    if not rows:
        return content
    md_rows = []
    for row_html in rows:
        cells = re.findall(
            r"<t[hd]\b[^>]*>(.*?)</t[hd]>", row_html, re.DOTALL | re.IGNORECASE
        )
        clean = [
            re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", cell)).strip() for cell in cells
        ]
        if clean:
            md_rows.append(clean)
    if not md_rows:
        return content
    col_count = max(len(r) for r in md_rows)
    lines = [
        "| " + " | ".join(md_rows[0]) + " |",
        "| " + " | ".join(["---"] * len(md_rows[0])) + " |",
    ]
    for row in md_rows[1:]:
        padded = row + [""] * (col_count - len(row))
        lines.append("| " + " | ".join(padded[:col_count]) + " |")
    md_table = "\n".join(lines)
    return (prefix + "\n\n" + md_table) if prefix else md_table


def _split_large_table(table: str) -> list[str]:
    rows = table.split("\n")
    sep_idx = next(
        (i for i, r in enumerate(rows) if re.match(r"^\|[\s\-|:]+\|", r)), None
    )
    if sep_idx is None or sep_idx == 0:
        return [table]
    header = "\n".join(rows[: sep_idx + 1])
    data_rows = rows[sep_idx + 1 :]
    chunks, current_rows = [], []
    current_size = len(header) + 1
    for row in data_rows:
        row_size = len(row) + 1
        if current_size + row_size > _TABLE_MAX_CHARS and current_rows:
            chunks.append(header + "\n" + "\n".join(current_rows))
            current_rows, current_size = [], len(header) + 1
        current_rows.append(row)
        current_size += row_size
    if current_rows:
        chunks.append(header + "\n" + "\n".join(current_rows))
    return chunks if chunks else [table]


def _segment_by_tables(md: str) -> list[tuple[str, str]]:
    lines = md.split("\n")
    segments, i = [], 0
    table_line_re = re.compile(r"^[|+]")
    html_table_re = re.compile(r"^<table(\s|>)", re.IGNORECASE)
    table_caption_re = re.compile(r"^Таблица\s+\d+")
    while i < len(lines):
        if table_line_re.match(lines[i]) or html_table_re.match(lines[i].lstrip()):
            table = []
            if html_table_re.match(lines[i].lstrip()):
                while i < len(lines):
                    table.append(lines[i])
                    if re.search(r"</table>", lines[i], re.IGNORECASE):
                        i += 1
                        break
                    i += 1
            else:
                while i < len(lines) and table_line_re.match(lines[i]):
                    table.append(lines[i])
                    i += 1
            if segments and segments[-1][0] == "text":
                prev_lines = segments[-1][1].split("\n")
                for j in range(len(prev_lines) - 1, -1, -1):
                    if prev_lines[j].strip():
                        if table_caption_re.match(prev_lines[j].strip()):
                            caption = prev_lines[j]
                            new_text = "\n".join(
                                prev_lines[:j] + prev_lines[j + 1 :]
                            ).rstrip()
                            segments[-1] = ("text", new_text)
                            table.insert(0, caption)
                        break
            segments.append(("table", "\n".join(table)))
        else:
            text = []
            while (
                i < len(lines)
                and not table_line_re.match(lines[i])
                and not html_table_re.match(lines[i].lstrip())
            ):
                text.append(lines[i])
                i += 1
            content = "\n".join(text)
            if content.strip():
                segments.append(("text", content))
    return segments


# ─── Captioning ──────────────────────────────────────────────────────────────


def _make_caption_fn(
    url: str, model: str, key: str
) -> Callable[[bytes, str, str], str]:
    client = OpenAI(base_url=url, api_key=key)

    def caption(img_bytes: bytes, mime: str, context: str) -> str:
        b64 = base64.b64encode(img_bytes).decode()
        prompt = (
            "Ты анализируешь изображение из технической документации. "
            "Ответь одним абзацем из 2–3 предложений: что изображено, "
            "какие основные элементы или обозначения видны. "
            "Не используй markdown — только обычный текст."
        )
        if context:
            prompt += f"\n\nКонтекст из документа: {context}"
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{mime};base64,{b64}"},
                            },
                        ],
                    }
                ],
                temperature=0.1,
                extra_body={"think": False},
                timeout=120,
            )
            text = resp.choices[0].message.content.strip()
            text = re.sub(r"\*{1,2}([^*]+)\*{1,2}", r"\1", text)
            text = re.sub(r"^#{1,4}\s+", "", text, flags=re.MULTILINE)
            text = re.sub(r"\n+", " ", text).strip()
            if len(text) > 400:
                cut = text[:400].rfind(".")
                text = text[: cut + 1] if cut > 0 else text[:400]
            return text
        except Exception as e:
            logger.warning(f"Captioning error: {e}")
            return ""

    return caption


# ─── DOCX: Floating Images ───────────────────────────────────────────────────


def _extract_floating_images(docx_path: Path) -> list[tuple[str, list[str]]]:
    from docx import Document

    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            rels_xml = z.read("word/_rels/document.xml.rels").decode()
        rels = dict(re.findall(r'Id="(rId\d+)"[^>]*Target="([^"]+)"', rels_xml))
        doc = Document(str(docx_path))
        result, last_nonempty = [], ""
        for para in doc.paragraphs:
            text, xml = para.text.strip(), para._p.xml
            if "wp:anchor" not in xml:
                if text:
                    last_nonempty = text
                continue
            rids = re.findall(r'r:embed="(rId\d+)"', xml)
            files = [
                rels[rid]
                for rid in rids
                if rid in rels and rels[rid].startswith("media/")
            ]
            files = [f for f in files if Path(f).suffix.lower() not in _SKIP_IMG_EXTS]
            if not files:
                if text:
                    last_nonempty = text
                continue
            anchor_text = text if text else last_nonempty
            if anchor_text:
                result.append((anchor_text, files))
            if text:
                last_nonempty = text
        return result
    except Exception as e:
        logger.warning(f"Floating images extraction failed: {e}")
        return []


def _inject_floating_images(md: str, floating: list[tuple[str, list[str]]]) -> str:
    occurrence_counter, insertions = {}, []
    for anchor_text, files in floating:
        key = re.sub(r"[^\w\s]", "", anchor_text[:40]).strip()
        if not key:
            continue
        target_occurrence = occurrence_counter.get(key, 0)
        occurrence_counter[key] = target_occurrence + 1
        pattern = re.compile(re.escape(key[:30]), re.IGNORECASE)
        matches = list(pattern.finditer(md))
        if not matches:
            continue
        if target_occurrence >= len(matches):
            target_occurrence = len(matches) - 1
        match_end = matches[target_occurrence].end()
        next_para = md.find("\n\n", match_end)
        insert_pos = next_para if next_para != -1 else len(md)
        imgs_md = "\n\n" + "\n\n".join(f"![]({f})" for f in files)
        insertions.append((insert_pos, imgs_md))
    for insert_pos, imgs_md in sorted(insertions, key=lambda x: x[0], reverse=True):
        md = md[:insert_pos] + imgs_md + md[insert_pos:]
    return md


# ─── Chunker Classes ─────────────────────────────────────────────────────────


class DoclingDocxChunker:
    def __init__(
        self,
        docx_path: str,
        document_name: str,
        images_dir: str = "docx_images",
        chunk_size: int = 1500,
        chunk_overlap: int = 150,
        debug_dir: str = "data/debug_chunks",
    ):
        self.docx_path = Path(docx_path)
        self.document_name = document_name
        safe = re.sub(r"[^\w\s\-\.]", "_", document_name).strip()
        self.safe_name = re.sub(r"[-\s]+", "_", safe)[:100]
        self.images_dir = Path(images_dir) / self.safe_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.debug_dir = Path(debug_dir)

        # Настройка Captioning
        self.caption_fn = None
        if os.getenv("CAPTION_MODEL"):
            self.caption_fn = _make_caption_fn(
                url=os.getenv("CAPTION_URL", "http://localhost:11434/v1"),
                model=os.getenv("CAPTION_MODEL", "qwen3.5:9b"),
                key=os.getenv("CAPTION_KEY", "ollama"),
            )

    def get_chunks(self) -> List[DocumentChunk]:
        # 1. Конвертация DOCX -> Markdown (Pandoc)
        md = self._convert_to_md()

        # 2. Извлечение таблиц отдельно (через python-docx для спец. коллекции)
        table_chunks = self._extract_table_chunks()

        # 3. Разбивка основного Markdown на чанки
        text_chunks = self._split_md_to_chunks(md)

        all_chunks = text_chunks + table_chunks
        self._save_debug(all_chunks)
        return all_chunks

    def _convert_to_md(self) -> str:
        logger.info(f"📄 Pandoc: {self.docx_path.name} ...")
        with tempfile.TemporaryDirectory() as tmpdir:
            src = self.docx_path
            if src.suffix.lower() == ".doc":
                subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        tmpdir,
                        str(src),
                    ],
                    check=True,
                    capture_output=True,
                )
                src = Path(tmpdir) / (src.stem + ".docx")
            raw = pypandoc.convert_file(str(src), "gfm", extra_args=["--wrap=none"])

        # Постобработка: чистка и инъекция картинок
        md = _strip_image_attrs(raw)
        floating = _extract_floating_images(self.docx_path)
        if floating:
            md = _inject_floating_images(md, floating)

        # Сохранение картинок на диск для мультимодальности
        self._save_images_to_disk()

        # Генерируем подписи (captioning)
        if self.caption_fn:
            md = self._apply_captions_to_md(md)

        return md

    def _save_images_to_disk(self):
        """Извлекает картинки из DOCX и сохраняет в images_dir."""
        self.images_dir.mkdir(parents=True, exist_ok=True)
        try:
            with zipfile.ZipFile(self.docx_path) as zf:
                for name in zf.namelist():
                    if (
                        name.startswith("word/media/")
                        and Path(name).suffix.lower() not in _SKIP_IMG_EXTS
                    ):
                        out_path = self.images_dir / Path(name).name
                        with open(out_path, "wb") as f:
                            f.write(zf.read(name))
        except Exception as e:
            logger.warning(f"Failed to save DOCX images: {e}")

    def _apply_captions_to_md(self, md: str) -> str:
        seen_captions = set()
        img_re = re.compile(r"!\[([^\]]*)\]\(((?:word/)?media/[^)]+)\)")

        with zipfile.ZipFile(self.docx_path) as zf:
            zip_names = set(zf.namelist())

            def replace_img(m):
                rel_path = m.group(2)
                zip_path = rel_path if rel_path in zip_names else f"word/{rel_path}"
                if zip_path not in zip_names:
                    return m.group(0)
                try:
                    img_bytes = zf.read(zip_path)
                    context = _extract_context(md, m.start(), m.end())
                    caption = self.caption_fn(img_bytes, "image/png", context)
                    if not caption:
                        return m.group(0)
                    sig = caption.lower()[:120]
                    if sig in seen_captions:
                        return f"![{_DUP_CAPTION_MARKER}]({rel_path})"
                    seen_captions.add(sig)
                    return f"![{caption}]({rel_path})"
                except:
                    return m.group(0)

            return img_re.sub(replace_img, md)

    def _extract_table_chunks(self) -> List[DocumentChunk]:
        """Использует логику из /evg/rag для точного извлечения таблиц."""
        try:
            doc = DocxDocument(str(self.docx_path))
            table_chunks = []
            last_text, table_idx = "", 0

            # Helper to check for table header flag
            def is_header_row(row):
                try:
                    return len(row._tr.get_or_add_trPr().xpath("./w:tblHeader")) > 0
                except:
                    return False

            # Iterate blocks
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            for child in doc.element.body.iterchildren():
                if isinstance(child, CT_P):
                    last_text = Paragraph(child, doc).text.strip() or last_text
                elif isinstance(child, CT_Tbl):
                    tbl = Table(child, doc)
                    table_idx += 1
                    # Extract header
                    header_rows, header_indices, sep = [], [], ""
                    for ri, row in enumerate(tbl.rows):
                        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                        header_rows.append("| " + " | ".join(cells) + " |")
                        header_indices.append(ri)
                        if is_header_row(row):
                            if ri + 1 < len(tbl.rows) and is_header_row(
                                tbl.rows[ri + 1]
                            ):
                                continue
                            sep = "| " + " | ".join(["---"] * len(cells)) + " |"
                            break
                        elif ri == 0:
                            sep = "| " + " | ".join(["---"] * len(cells)) + " |"
                            break

                    full_header = "\n".join(header_rows + [sep])
                    context = (
                        f"{last_text}\n\n" if last_text and len(last_text) < 300 else ""
                    )
                    prefix = f"Таблица из '{self.document_name}':\n\n{context}"

                    # Split data rows
                    data_indices = [
                        i for i in range(len(tbl.rows)) if i not in header_indices
                    ]
                    current_rows, current_len, sub_idx = (
                        [],
                        len(prefix) + len(full_header) + 2,
                        1,
                    )

                    def flush(rows, sidx):
                        if not rows:
                            return
                        content = prefix + full_header + "\n" + "\n".join(rows)
                        table_chunks.append(
                            DocumentChunk(
                                content=content,
                                metadata={
                                    "document_name": self.document_name,
                                    "chunk_index": table_idx,
                                    "sub_chunk_index": sidx,
                                    "is_table": True,
                                    "table_caption": last_text if context else "",
                                },
                                source_type="docx_table",
                                chunk_id=f"{self.document_name}_table_{table_idx}_{sidx}",
                            )
                        )

                    for ri in data_indices:
                        row_md = (
                            "| "
                            + " | ".join(
                                [
                                    c.text.strip().replace("\n", " ")
                                    for c in tbl.rows[ri].cells
                                ]
                            )
                            + " |"
                        )
                        if (
                            current_len + len(row_md) + 1 > self.chunk_size
                            and current_rows
                        ):
                            flush(current_rows, sub_idx)
                            current_rows, current_len, sub_idx = (
                                [row_md],
                                len(prefix) + len(full_header) + len(row_md) + 2,
                                sub_idx + 1,
                            )
                        else:
                            current_rows.append(row_md)
                            current_len += len(row_md) + 1
                    flush(current_rows, sub_idx)
                    last_text = ""
            return table_chunks
        except Exception as e:
            logger.error(f"Table extraction failed: {e}")
            return []

    def _split_md_to_chunks(self, md: str) -> List[DocumentChunk]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
        )
        texts = []
        for seg_type, content in _segment_by_tables(md):
            if seg_type == "table":
                if re.search(r"<table\b", content, re.IGNORECASE):
                    content = _html_table_to_md(content)
                if len(content) > _TABLE_MAX_CHARS:
                    texts.extend(_split_large_table(content))
                else:
                    texts.append(content)
            else:
                texts.extend(splitter.split_text(content))

        # Извлекаем пути к картинкам для каждого чанка (для мультимодальности)
        img_pattern = re.compile(r"!\[.*?\]\(((?:word/)?media/[^)]+)\)")

        chunks = []
        for i, text in enumerate(texts):
            # Мапим пути картинок на реальные пути на диске
            found_imgs = img_pattern.findall(text)
            real_paths = [str(self.images_dir / Path(p).name) for p in found_imgs]

            chunks.append(
                DocumentChunk(
                    content=text,
                    metadata={
                        "document_name": self.document_name,
                        "chunk_index": i + 1,
                        "total_chunks": len(texts),
                        "image_paths": real_paths,
                    },
                    source_type="docx",
                    chunk_id=f"{self.document_name}_chunk_{i + 1}",
                )
            )
        return chunks

    def _save_debug(self, chunks: List[DocumentChunk]):
        self.debug_dir.mkdir(parents=True, exist_ok=True)
        with open(
            self.debug_dir / f"{self.safe_name}_chunks.md", "w", encoding="utf-8"
        ) as f:
            for ch in chunks:
                f.write(
                    f"{'=' * 80}\n### {ch.source_type.upper()} {ch.metadata.get('chunk_index')}\nID: {ch.chunk_id}\n{'-' * 80}\n\n{ch.content}\n\n"
                )


class PyMuPDF4LLMPDFChunker(DoclingDocxChunker):
    def __init__(
        self,
        pdf_path: str,
        document_name: str,
        images_dir: str = "data/pdf_images",
        chunk_size: int = 1500,
        chunk_overlap: int = 150,
        debug_dir: str = "data/debug_chunks",
    ):
        super().__init__(
            pdf_path, document_name, images_dir, chunk_size, chunk_overlap, debug_dir
        )
        self.pdf_path = Path(pdf_path)
        self.glm_url = os.getenv("GLM_OCR_URL", "http://localhost:11434/v1")
        self.glm_model = os.getenv("GLM_OCR_MODEL", "glm-ocr:latest")
        self.glm_key = os.getenv("GLM_OCR_KEY", "ollama")

    def get_chunks(self) -> List[DocumentChunk]:
        md, used_ocr = self._convert()
        source_type = "pdf_ocr" if used_ocr else "pdf"

        # Таблицы (из Docling или OCR)
        table_chunks = self._extract_pdf_table_chunks(md, used_ocr)

        # Основные чанки
        text_chunks = self._split_pdf_md_to_chunks(md, source_type)

        all_chunks = text_chunks + table_chunks
        self._save_debug(all_chunks)
        return all_chunks

    def _convert(self) -> Tuple[str, bool]:
        logger.info(f"📄 Docling: {self.pdf_path.name} ...")
        try:
            with fitz.open(str(self.pdf_path)) as doc:
                num_pages = len(doc)

            pipeline_options = PdfPipelineOptions(do_ocr=False, do_table_structure=True)
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
            result = converter.convert(str(self.pdf_path), raises_on_error=True)
            self._docling_result = result

            md = result.document.export_to_markdown()
            # Captioning & Image saving
            if self.caption_fn:
                md = self._apply_pdf_captions(result)

            md = re.sub(r"\n{3,}", "\n\n", md).strip()
            if is_garbage(md, 0.30, num_pages=num_pages):
                logger.warning("   ⚠️ Garbage detected, falling back to GLM OCR")
                return self._glm_fallback(), True
            return md, False
        except Exception as e:
            logger.error(f"Docling failed: {e}, using GLM OCR fallback")
            return self._glm_fallback(), True

    def _glm_fallback(self) -> str:
        import glmocr
        from urllib.parse import urlparse

        p = urlparse(self.glm_url)
        logger.info(f"   GLM OCR: {self.pdf_path.name} ...")
        with glmocr.GlmOcr(
            mode="selfhosted",
            ocr_api_host=p.hostname or "127.0.0.1",
            ocr_api_port=p.port or 11434,
            log_level="ERROR",
        ) as parser:
            res = parser.parse(str(self.pdf_path))
        md = res.markdown_result or ""
        # Сохранение картинок из OCR (glmocr сохраняет их в imgs/)
        return re.sub(r"\n{3,}", "\n\n", md).strip()

    def _apply_pdf_captions(self, result) -> str:
        md = result.document.export_to_markdown()
        pictures = list(result.document.pictures)
        if not pictures:
            return md
        self.images_dir.mkdir(parents=True, exist_ok=True)
        placeholder, doc = "<!-- image -->", fitz.open(str(self.pdf_path))
        ph_positions = [m.start() for m in re.finditer(re.escape(placeholder), md)]
        seen_captions = set()
        for i, pic in enumerate(pictures):
            if not pic.prov or i >= len(ph_positions):
                continue
            prov = pic.prov[0]
            p = doc[prov.page_no - 1]
            b = prov.bbox
            clip = fitz.Rect(b.l, p.rect.height - b.t, b.r, p.rect.height - b.b)
            try:
                img_bytes = p.get_pixmap(clip=clip, dpi=150).tobytes("png")
                filename = f"pic_{i + 1}.png"
                (self.images_dir / filename).write_bytes(img_bytes)
                context = _extract_context(
                    md, ph_positions[i], ph_positions[i] + len(placeholder)
                )
                caption = self.caption_fn(img_bytes, "image/png", context)
                if not caption:
                    continue
                sig = caption.lower()[:120]
                text = _DUP_CAPTION_MARKER if sig in seen_captions else caption
                seen_captions.add(sig)
                md = md.replace(placeholder, f"![{text}](imgs/{filename})", 1)
            except:
                pass
        doc.close()
        return md

    def _extract_pdf_table_chunks(self, md: str, used_ocr: bool) -> List[DocumentChunk]:
        table_chunks = []
        if used_ocr:
            # Парсим HTML таблицы из OCR
            for m in re.finditer(r"<table\b[^>]*>.*?</table>", md, re.DOTALL):
                table_md = _html_table_to_md(m.group(0))
                if not re.sub(r"[|\-\s]", "", table_md):
                    continue
                table_chunks.append(
                    DocumentChunk(
                        content=table_md,
                        metadata={
                            "document_name": self.document_name,
                            "is_table": True,
                        },
                        source_type="pdf_table",
                        chunk_id=f"{self.document_name}_table_{len(table_chunks) + 1}",
                    )
                )
        else:
            # Таблицы из Docling
            if hasattr(self, "_docling_result"):
                doc = self._docling_result.document
                items = list(doc.iterate_items())
                for i, (item, _) in enumerate(items):
                    if isinstance(item, TableItem):
                        t_md = item.export_to_markdown(doc=doc).strip()
                        if not t_md:
                            continue
                        table_chunks.append(
                            DocumentChunk(
                                content=t_md,
                                metadata={
                                    "document_name": self.document_name,
                                    "is_table": True,
                                },
                                source_type="pdf_table",
                                chunk_id=f"{self.document_name}_table_{len(table_chunks) + 1}",
                            )
                        )
        return table_chunks

    def _split_pdf_md_to_chunks(self, md: str, source_type: str) -> List[DocumentChunk]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["<!-- page ", "\n\n", "\n", " ", ""],
        )
        texts = splitter.split_text(md)
        img_pattern = re.compile(r"!\[.*?\]\((imgs/[^)]+)\)")
        chunks = []
        for i, text in enumerate(texts):
            found_imgs = img_pattern.findall(text)
            real_paths = [str(self.images_dir / Path(p).name) for p in found_imgs]
            chunks.append(
                DocumentChunk(
                    content=text,
                    metadata={
                        "document_name": self.document_name,
                        "chunk_index": i + 1,
                        "total_chunks": len(texts),
                        "image_paths": real_paths,
                    },
                    source_type=source_type,
                    chunk_id=f"{self.document_name}_chunk_{i + 1}",
                )
            )
        return chunks
